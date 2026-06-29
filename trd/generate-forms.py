#!/usr/bin/env python3
"""Generate TRD Excel and Word forms for AI Red Teaming.

Produces 4 files in red-teaming/:
  trd-red-teaming-consultant.xlsx  — Excel with guidance + examples
  trd-red-teaming-customer.xlsx    — Excel clean
  trd-red-teaming-consultant.docx  — Word with guidance + examples
  trd-red-teaming-customer.docx    — Word clean

Usage:
  python3 generate-forms.py
"""

import os
from copy import copy

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT_DIR = os.path.join(os.path.dirname(__file__), "red-teaming")

# ── Colors ────────────────────────────────────────────────────────
PAN_NAVY = "00294D"
PAN_GOLD = "FFB81C"
WHITE = "FFFFFF"
FILL_YELLOW = "FFF9E6"
FILL_PURPLE = "F3E8FF"
LIGHT_GRAY = "F7F8FA"
BORDER_GRAY = "DEE2E6"
RED = "EF4444"
AMBER = "F59E0B"


# ═══════════════════════════════════════════════════════════════════
# TRD DATA — single source of truth for all generated forms
# ═══════════════════════════════════════════════════════════════════

# Q = question dict: marker, text, fmt, options, consultant_note
def Q(text, fmt="text", options=None, marker="", note=""):
    return {"marker": marker, "text": text, "fmt": fmt,
            "options": options or [], "note": note}

# Contact row helper
def CONTACT(role, marker="*"):
    return {"role": role, "marker": marker}


PARTS = [
    # ── PART A ────────────────────────────────────────────────────
    {
        "part": "Part A — Engagement Context",
        "desc": "Asked once per customer",
        "tab_name": "Part A — Engagement",
        "sections": [
            {
                "num": "1", "title": "Engagement Overview", "tier": "Scoping",
                "feeds": "",
                "questions": [
                    Q("Customer name", "text", marker="*"),
                    Q("Primary business driver for AI red teaming?", "select",
                      ["Pre-deployment security validation", "Compliance requirement",
                       "Risk assessment", "Incident response", "Other"], "*"),
                    Q("What does success look like for this engagement?", "text", marker="*"),
                    Q("Target start date", "date", marker="*"),
                    Q("Desired completion date", "date"),
                ],
                "contacts": [
                    CONTACT("Project sponsor", "*"),
                    CONTACT("Technical lead (AI/ML)", "*"),
                    CONTACT("Security lead", "*"),
                    CONTACT("CSP portal admin", ""),
                ],
                "consultant_note": (
                    "The business driver shapes scope. Compliance-driven engagements need "
                    "compliance-mapped reports (OWASP, NIST). Risk assessment prioritizes "
                    "breadth across targets.\n\n"
                    "Common pattern: Most customers cite 'pre-deployment validation' but actually "
                    "need compliance evidence for auditors. Probe for the real audience of the "
                    "final report.\n\n"
                    "If they don't know: Default to 'Pre-deployment security validation' and "
                    "adjust after the first scan cycle."
                ),
            },
            {
                "num": "2", "title": "Target Overview", "tier": "Essential",
                "feeds": "Drives Part C count",
                "questions": [],
                "target_table": True,
                "consultant_note": (
                    "Row count determines Part C sections, credit consumption, and package.\n\n"
                    "Target types:\n"
                    "• Model — Direct model endpoint (simplest config)\n"
                    "• Application — AI-powered app with business logic, guardrails (needs "
                    "industry, use case, competitors)\n"
                    "• Agent — Autonomous with tool access (needs tool schemas for "
                    "tool-misuse attacks)\n\n"
                    "Gotcha: A single model behind 3 apps = 3 targets (different system "
                    "prompts, guardrails, use cases)."
                ),
            },
            {
                "num": "3", "title": "Compliance Requirements", "tier": "Important",
                "feeds": "Step 6.2 (category selection)",
                "questions": [
                    Q("Which compliance frameworks must scan results map to?", "multi-select",
                      ["OWASP LLM Top 10", "MITRE ATLAS", "NIST AI RMF",
                       "DASF V2.0", "None specific", "Other"]),
                    Q("Are compliance-mapped reports required for auditors or regulators?", "select",
                      ["Yes — auditor delivery", "Yes — internal compliance", "No"]),
                    Q("Industry-specific AI regulations that apply?", "text",
                      note="e.g., EU AI Act, FDA AI/ML guidance, FFIEC"),
                ],
                "consultant_note": (
                    "Framework selection drives COMPLIANCE subcategories in scan config.\n\n"
                    "Common: Financial services → NIST + OWASP. Healthcare → NIST. "
                    "Tech → OWASP + MITRE ATLAS. Government → NIST mandatory.\n\n"
                    "If they don't know: Enable all four. No performance penalty."
                ),
            },
            {
                "num": "4", "title": "Remediation Expectations", "tier": "Important",
                "feeds": "Step 8.3 (remediation review)",
                "questions": [
                    Q("Will Red Teaming findings feed into AIRS Runtime for real-time protection?",
                      "select",
                      ["Yes — RT→Runtime bridge planned", "Under evaluation",
                       "No — findings for manual remediation only"]),
                    Q("Who will implement remediation for discovered vulnerabilities?", "select",
                      ["Customer AI/ML team", "Customer security team",
                       "Palo Alto Networks PS", "Joint effort"]),
                ],
                "consultant_note": (
                    "RT→Runtime bridge is a key upsell. If planned, coordinate deployment "
                    "profiles and ensure Runtime API Intercept is licensed alongside Red Teaming.\n\n"
                    "Customers who say 'manual remediation only' often change their mind after "
                    "seeing the first report."
                ),
            },
        ],
    },

    # ── PART B ────────────────────────────────────────────────────
    {
        "part": "Part B — Platform Configuration",
        "desc": "Asked once per engagement",
        "tab_name": "Part B — Platform",
        "sections": [
            {
                "num": "5", "title": "Licensing & Credits", "tier": "Essential",
                "feeds": "Step 3.1 (Create Deployment Profile)",
                "questions": [
                    Q("CSP admin email (with credit allocation permissions)", "email", marker="*"),
                    Q("Total credits in customer's AIRS pool?", "number", marker="*"),
                    Q("Credits allocated specifically for AI Red Teaming?", "number", marker="*"),
                    Q("Are credits shared with other AIRS products (Runtime, Model Security)?",
                      "select", ["Yes — shared pool", "No — dedicated RT allocation"]),
                ],
                "consultant_note": (
                    "Credits determine scan quota. A single Attack Library scan ≈ 1 session "
                    "per target.\n\n"
                    "Gotcha: CSP admin must have the credit allocation role — standard SCM "
                    "admins cannot allocate credits. Verify during the pre-engagement call."
                ),
            },
            {
                "num": "6", "title": "Tenant & TSG", "tier": "Essential",
                "feeds": "Step 3.1 (Create Deployment Profile)",
                "questions": [
                    Q("Does a Tenant Service Group (TSG) already exist, or will one be created?",
                      "select", ["Existing TSG", "Create new TSG"], "*"),
                    Q("If existing: TSG ID", "text", marker="†"),
                    Q("Preferred region for the AIRS deployment profile?", "select",
                      ["Americas", "EU", "Singapore"], "*"),
                    Q("Is Strata Logging Service (SLS) enabled on this tenant?", "select",
                      ["Yes", "No", "Unknown"], "*"),
                    Q("Does the tenant have an existing AIOps subscription?", "select",
                      ["Yes", "No", "Unknown"]),
                ],
                "consultant_note": (
                    "TSG and region cannot be changed after activation. SLS is mandatory — "
                    "activate first if not enabled (adds 10–15 min).\n\n"
                    "Existing AIOps can conflict with AIRS onboarding. New TSG adds 15–20 "
                    "min provisioning."
                ),
            },
            {
                "num": "7", "title": "Access & Identity", "tier": "Essential",
                "feeds": "Step 3.2 (Configure IAM)",
                "questions": [
                    Q("Which RBAC roles are needed for the engagement?", "multi-select",
                      ["Superuser", 'Custom role with "AI Red Teaming" enabled',
                       "Read-only for stakeholders"], "*"),
                    Q("Is a service account needed for API or CI/CD access?", "select",
                      ["Yes", "No", "TBD"], "*"),
                    Q("Does the customer use SSO / IdP for SCM access?", "select",
                      ["Okta", "Entra ID (Azure AD)", "Ping Identity",
                       "Google Workspace", "No SSO", "Other"]),
                    Q("How many users need SCM access for this engagement?", "number"),
                ],
                "consultant_note": (
                    "Custom roles must explicitly enable 'AI Red Teaming' app access. "
                    "Service accounts required for API scans — record Client Secret "
                    "immediately (not retrievable later).\n\n"
                    "Most engagements need: 1 Superuser (PS) + 1 read-only (stakeholders) "
                    "+ 1 service account (API)."
                ),
            },
            {
                "num": "8", "title": "Network & Connectivity", "tier": "Essential",
                "feeds": "Phase 5 decision (skip or deploy Network Channel)",
                "questions": [
                    Q("Are the target endpoints publicly accessible from the internet?", "select",
                      ["All public", "All private", "Mix of public and private"], "*"),
                    Q("If private: is a Kubernetes cluster available for the Network Channel client?",
                      "select",
                      ["Yes — managed K8s (EKS/AKS/GKE)", "Yes — self-managed K8s",
                       "No K8s available"], "†"),
                    Q("If K8s: does the cluster have outbound internet access?", "select",
                      ["Yes", "No — air-gapped", "Restricted (proxy/allowlist)"], "†"),
                    Q("Is there a WAF or API gateway in the request path to any target?", "select",
                      ["Yes", "No", "Unknown"]),
                    Q("If WAF present: desired testing strategy?", "select",
                      ["Whitelist AIRS source IPs", "Bypass WAF (test app directly)",
                       "Test full stack including WAF"], "†"),
                ],
                "consultant_note": (
                    "Single most impactful scoping question. Public → skip Phase 5 entirely. "
                    "Private → deploy Network Channel (1–2 hours + K8s prereqs).\n\n"
                    "WAF decision: Testing through WAF = realistic full-stack test. Bypassing "
                    "WAF = isolates AI-specific vulnerabilities. Neither is wrong — clarify intent."
                ),
            },
            {
                "num": "9", "title": "Engagement Sizing", "tier": "Scoping",
                "feeds": "Calculated from §2, §5, §8",
                "questions": [
                    Q("Selected package", "select",
                      ["Starter (1 target)", "Standard (up to 3)",
                       "Advanced (up to 5 + Network Channel)",
                       "Enterprise (5+ custom)"]),
                    Q("Add-on: Network Channel deployment?", "select", ["Yes", "No"]),
                    Q("Add-on: REST wrapper development?", "select", ["Yes", "No"]),
                    Q("Add-on: CI/CD pipeline integration?", "select", ["Yes", "No"]),
                    Q("Add-on: Custom attack development?", "select", ["Yes", "No"]),
                ],
                "consultant_note": (
                    "Sizing formula: (targets × scan types × scan cycles) = total sessions.\n"
                    "Attack Library ≈ 5 hours/target. Dynamic Agent ≈ 1–2 hours.\n\n"
                    "Pre-built connections (OpenAI, Bedrock, Vertex) < 10 min to configure. "
                    "REST (custom) connections: budget 30–60 min per target for template work."
                ),
            },
        ],
    },

    # ── PART C ────────────────────────────────────────────────────
    {
        "part": "Part C — Per-Target Configuration",
        "desc": "Repeated for each target. Duplicate this tab for additional targets.",
        "tab_name": "Part C — Target 1",
        "sections": [
            {
                "num": "10", "title": "Target Architecture", "tier": "Essential",
                "feeds": "Step 4.2 (Add Target) + Step 4.3 (Configure Details)",
                "questions": [
                    Q("Application name (human-readable)", "text", marker="*"),
                    Q("Target type", "select", ["MODEL", "APPLICATION", "AGENT"], "*"),
                    Q("Inference provider", "select",
                      ["OPENAI", "AZURE_OPENAI", "BEDROCK", "VERTEX",
                       "HUGGINGFACE", "REST (custom)"], "*"),
                    Q("Base model", "text", marker="*",
                      note="e.g., GPT-4o, Claude 3.5, Gemini 2.0, Llama 3.1"),
                    Q("Core architecture", "select",
                      ["Single LLM", "RAG (retrieval-augmented)",
                       "Tool Calling", "Multi-Agent", "Other"]),
                    Q("Industry (required for Application and Agent types)", "text", marker="*",
                      note="e.g., Financial services, Healthcare, Technology"),
                    Q("Use case description", "text", marker="*",
                      note="e.g., Customer support chatbot for insurance claims"),
                    Q("Competitor names (for brand reputation attacks)", "text"),
                    Q("System prompt (enables white-box testing)", "text",
                      note="Paste or attach the full system prompt"),
                    Q("Tools accessible (Agent targets only — provide tool schemas)", "json",
                      marker="†"),
                ],
                "consultant_note": (
                    "System prompt is the single highest-impact optional field — enables "
                    "white-box testing, dramatically improves Agent scan results. If customer "
                    "is reluctant, explain it stays within the AIRS platform.\n\n"
                    "Agent tool schemas enable tool-misuse attacks — high-impact for agentic "
                    "architectures."
                ),
            },
            {
                "num": "11", "title": "Risk Priority Matrix", "tier": "Important",
                "feeds": "Step 6.2 (Configure Scan Categories)",
                "questions": [],
                "risk_matrix": True,
                "consultant_note": (
                    "Default: all Security → High, all Safety → Medium. Compliance per §3. "
                    "Brand → High for customer-facing apps, N/A for internal tools.\n\n"
                    "Don't set everything to High — makes reports less actionable. "
                    "Prioritization helps focus remediation."
                ),
            },
            {
                "num": "12", "title": "Connection & Auth", "tier": "Essential",
                "feeds": "Step 4.2 (Add Target)",
                "questions": [
                    Q("Connection type", "select",
                      ["OPENAI", "AZURE_OPENAI", "BEDROCK", "VERTEX",
                       "HUGGINGFACE", "REST"], "*"),
                    Q("Endpoint URL", "url", marker="*"),
                    Q("Endpoint accessibility", "select",
                      ["PUBLIC", "PRIVATE", "NETWORK_BROKER (API-only)"], "*"),
                    Q("Authentication method", "select",
                      ["HEADERS (API key)", "BASIC_AUTH", "OAUTH2"], "*"),
                    Q("Auth credentials (API key, token, or OAuth config)", "text", marker="*",
                      note="Provide securely — do not include in shared documents"),
                    Q("Token/credential expiry period?", "text",
                      note="e.g., 1 hour, 24 hours, no expiry"),
                    Q("API protocol", "select",
                      ["REST (default)", "WebSocket", "gRPC", "Socket.IO", "GraphQL"]),
                ],
                "consultant_note": (
                    "Token expiry > 1hr critical — Attack Library scans run ~5 hours. Static "
                    "API keys preferred for test envs.\n\n"
                    "NETWORK_BROKER is API-only (not in SCM UI). Non-REST protocols trigger "
                    "§16 (REST Wrapper)."
                ),
            },
            {
                "num": "13", "title": "Rate Limits & Capacity", "tier": "Essential",
                "feeds": "Step 4.3 + scan execution",
                "questions": [
                    Q("Requests per minute (RPM) limit?", "number", marker="*"),
                    Q("Tokens per minute (TPM) limit?", "number", marker="*"),
                    Q("Tokens per day (TPD) limit?", "number"),
                    Q("Maximum concurrent requests?", "number"),
                    Q("Maximum input tokens per request?", "number"),
                    Q("Maximum output tokens per request?", "number"),
                ],
                "consultant_note": (
                    "Rate limits are the #1 cause of scan failures. Attack Library needs "
                    "10–20 RPM sustained for ~5 hours.\n\n"
                    "Common:\n"
                    "• OpenAI Tier 1: 500 RPM, 30K TPM — fine\n"
                    "• OpenAI Tier 0/free: 3 RPM — too low, will fail\n"
                    "• Bedrock: varies by model — check AWS console\n\n"
                    "Recommend a dedicated API key with elevated limits for testing."
                ),
            },
            {
                "num": "14", "title": "Guardrail Detection", "tier": "Essential",
                "feeds": "Step 4.3 (Configure Details)",
                "questions": [
                    Q("Does the target have guardrails (content filters, safety layers)?",
                      "select", ["Yes", "No", "Unknown"], "*"),
                    Q("HTTP status code returned when a guardrail blocks a request?", "text",
                      marker="†", note="e.g., 400, 403, 200 (with error in body)"),
                    Q("Sample error response JSON when a guardrail blocks", "json", marker="†"),
                    Q("Guardrail technology in use?", "text",
                      note="e.g., Azure Content Safety, OpenAI moderation, custom, AIRS Runtime"),
                ],
                "consultant_note": (
                    "Without guardrail detection, reports can't distinguish 'blocked by "
                    "guardrail' from 'model refused' — makes reports nearly useless.\n\n"
                    "Common: OpenAI → 400 + content_policy_violation. Azure OpenAI → 200 + "
                    "content_filter in finish_reason.\n\n"
                    "If they don't know: Send a known-blocked prompt (e.g., 'How do I make "
                    "a bomb?') and capture the HTTP response."
                ),
            },
            {
                "num": "15", "title": "Environment", "tier": "Important",
                "feeds": "Result validity",
                "questions": [
                    Q("Which environment will be tested?", "select",
                      ["Production", "Staging", "Development", "Dedicated test"], "*"),
                    Q("Does this environment use the same model and config as production?",
                      "select",
                      ["Yes — identical", "Mostly — same model, different guardrails",
                       "No — different model or config"]),
                    Q("Is scanning production directly acceptable?", "select",
                      ["Yes", "No", "Only during maintenance windows"]),
                    Q("Maintenance window (if applicable)?", "text",
                      note="e.g., Saturdays 2am–6am UTC"),
                ],
                "consultant_note": (
                    "Non-production with different guardrails/models may not reflect production "
                    "risk. Document the delta for stakeholders.\n\n"
                    "Agent targets with tool access can trigger side effects in production — "
                    "clarify risk tolerance for agentic targets specifically."
                ),
            },
            {
                "num": "16", "title": "REST Wrapper (if non-REST protocol)", "tier": "Essential",
                "feeds": "Conditional: only if non-REST protocol in §12",
                "questions": [
                    Q("Native protocol of the target", "select",
                      ["WebSocket", "gRPC", "Socket.IO", "GraphQL", "Other"], "*"),
                    Q("Preferred language for the wrapper", "select",
                      ["Python", "Node.js", "Go", "Java", "Customer provides"], "*"),
                    Q("Does the target use streaming responses?", "select",
                      ["Yes — SSE", "Yes — chunked", "No — single response"]),
                    Q("If streaming: how should chunks be aggregated?", "select",
                      ["Concatenate all chunks", "Use last chunk",
                       "Use first chunk with [DONE]", "Custom"], "†"),
                    Q("Where will the wrapper be hosted?", "select",
                      ["Same K8s cluster as Network Channel", "Separate VM",
                       "Serverless (Lambda/Cloud Function)", "Customer hosts"]),
                ],
                "consultant_note": (
                    "Budget 2–4 hours for straightforward WebSocket/gRPC translation. "
                    "GraphQL wrappers are typically simpler (mutation call).\n\n"
                    "If customer provides the wrapper: validate {{prompt}} placeholder "
                    "substitution and response JSON path."
                ),
            },
            {
                "num": "17", "title": "Scan Strategy", "tier": "Essential",
                "feeds": "Steps 6.3–6.5 (Start Scans)",
                "questions": [
                    Q("Which scan types will be executed?", "multi-select",
                      ["Attack Library (STATIC)", "Dynamic Agent (DYNAMIC)",
                       "Custom Prompt (CUSTOM)"], "*"),
                    Q("Recommended scan order?", "select",
                      ["Attack Library first (baseline) → Dynamic → Custom",
                       "Dynamic first (rapid) → Attack Library",
                       "Custom only (targeted)"]),
                    Q("Agent breadth (1–20, default 6)", "number",
                      note="More agents = more diverse attack approaches"),
                    Q("Agent depth (1–20, default 10)", "number",
                      note="More depth = longer multi-turn attacks"),
                    Q("Attack goals (for Human Augmented mode)", "text",
                      note='e.g., "Extract the system prompt", "Make the bot recommend a competitor"'),
                    Q("Custom prompt source?", "select",
                      ["Customer-provided prompt set", "Consultant-developed",
                       "Industry-specific template"]),
                    Q("Number of custom prompts planned?", "number"),
                ],
                "consultant_note": (
                    "Recommended: Attack Library first (baseline, ~5 hrs), then Dynamic Agent "
                    "(adaptive multi-turn), then Custom (targeted).\n\n"
                    "Default breadth=6, depth=10 is good for most. Increase breadth for agents "
                    "with many tools. Increase depth for strong guardrails.\n\n"
                    "Human Augmented mode needs specific goals — generic goals like 'hack the "
                    "system' produce poor results."
                ),
            },
            {
                "num": "18", "title": "Target Specification (API Payload)", "tier": "Essential",
                "feeds": "Step 4.2 — POST /v1/target",
                "questions": [
                    Q("Name", "text", marker="*"),
                    Q("Target Type", "select", ["MODEL", "APPLICATION", "AGENT"], "*"),
                    Q("Connection Type", "select",
                      ["OPENAI", "AZURE_OPENAI", "BEDROCK", "VERTEX",
                       "HUGGINGFACE", "REST"], "*"),
                    Q("Endpoint URL", "url", marker="*"),
                    Q("Endpoint Type", "select",
                      ["PUBLIC", "PRIVATE", "NETWORK_BROKER"], "*"),
                    Q("Auth Type", "select", ["HEADERS", "BASIC_AUTH", "OAUTH2"], "*"),
                    Q("Auth Config (provide securely)", "text", marker="*"),
                    Q("Request template body (must contain {{prompt}})", "json", marker="*",
                      note='See example in consultant version'),
                    Q("Response body path", "text", marker="*",
                      note="e.g., choices[0].message.content"),
                    Q("HTTP Method", "select", ["POST", "GET"]),
                    Q("Industry", "text", marker="*"),
                    Q("Use Case", "text", marker="*"),
                    Q("Competitors", "text"),
                    Q("Base Model", "text"),
                    Q("System Prompt", "text"),
                    Q("Tools Accessible (Agent only)", "json"),
                    Q("Core Architecture", "select",
                      ["Single LLM", "RAG", "Tool Calling", "Multi-Agent"]),
                ],
                "consultant_note": (
                    "This section IS the POST /v1/target API payload. Validate before scanning:\n"
                    "1. Send test request with benign prompt — verify response\n"
                    "2. Verify response body path extracts actual model text\n"
                    "3. Test guardrail detection pattern\n"
                    "4. Confirm rate limits not hit\n\n"
                    "Request template example (OpenAI):\n"
                    '{\n'
                    '  "model": "gpt-4o",\n'
                    '  "messages": [\n'
                    '    {"role": "user", "content": "{{prompt}}"}\n'
                    '  ],\n'
                    '  "temperature": 0.7,\n'
                    '  "max_tokens": 1024\n'
                    '}\n\n'
                    "Common mistake: Azure OpenAI endpoint must include deployment name:\n"
                    "https://{resource}.openai.azure.com/openai/deployments/{deployment}/"
                    "chat/completions?api-version=2024-02-01"
                ),
            },
        ],
    },

    # ── DAY-2 ─────────────────────────────────────────────────────
    {
        "part": "Day-2 Appendix",
        "desc": "Optional — address after first scan cycle",
        "tab_name": "Day-2 Appendix",
        "sections": [
            {
                "num": "D1", "title": "Logging & SIEM Integration", "tier": "Day-2",
                "feeds": "",
                "questions": [
                    Q("SIEM platform in use?", "select",
                      ["Cortex XSIAM", "Splunk", "Sentinel", "QRadar", "Elastic", "Other"]),
                    Q("Should AIRS scan events forward to SIEM?", "select",
                      ["Yes", "No", "TBD"]),
                    Q("Log forwarding method preference?", "select",
                      ["Syslog", "API webhook",
                       "Strata Logging Service → SIEM connector", "Manual export"]),
                    Q("Alert thresholds?", "text",
                      note="e.g., Any Critical finding, Risk score > 7"),
                    Q("Log retention requirements?", "text",
                      note="e.g., 90 days, 1 year, regulatory requirement"),
                ],
                "consultant_note": "",
            },
            {
                "num": "D2", "title": "Ongoing Program Design", "tier": "Day-2",
                "feeds": "",
                "questions": [
                    Q("Desired scan cadence after initial engagement?", "select",
                      ["Weekly", "Monthly", "Quarterly", "On-demand only",
                       "Before each release"]),
                    Q("Should scans integrate into CI/CD pipeline?", "select",
                      ["Yes — gating (block deploy)", "Yes — advisory (report only)", "No"]),
                    Q("CI/CD platform in use?", "select",
                      ["GitHub Actions", "GitLab CI", "Jenkins", "Azure DevOps", "Other"]),
                    Q("Risk score threshold for automated re-scan trigger?", "number",
                      note="e.g., re-scan if overall risk > 5"),
                    Q("Go/no-go risk thresholds for deployment decisions?", "text",
                      note="e.g., No Critical findings, overall risk ≤ 3"),
                ],
                "consultant_note": "",
            },
            {
                "num": "D3", "title": "Reporting & Distribution", "tier": "Day-2",
                "feeds": "",
                "questions": [
                    Q("Report recipients (beyond engagement team)?", "text",
                      note="Name, email, report type (full / executive summary)"),
                    Q("Is an executive summary report required?", "select", ["Yes", "No"]),
                    Q("Report format preference?", "select",
                      ["PDF", "HTML", "JSON (API)", "All formats"]),
                    Q("Should reports be retained as audit evidence?", "select",
                      ["Yes — regulatory", "Yes — internal policy", "No"]),
                ],
                "consultant_note": "",
            },
        ],
    },
]

RISK_CATEGORIES = {
    "Security": [
        "Adversarial Suffix", "Evasion", "Indirect Prompt Injection",
        "Jailbreak", "Multi-Turn", "Prompt Injection",
        "Remote Code Execution", "System Prompt Leak", "Tool Leak",
    ],
    "Safety": [
        "Bias", "CBRN", "Cybercrime", "Drugs", "Non-Violent Crimes",
        "Political", "Self-Harm", "Sexual", "Violent Crimes & Weapons",
    ],
    "Compliance": [
        "OWASP LLM Top 10", "MITRE ATLAS", "NIST AI RMF", "DASF V2.0",
    ],
    "Brand": [
        "Off-Brand Responses", "Competitor Endorsement",
        "Product Misinformation", "Tone Inconsistency",
    ],
}


# ═══════════════════════════════════════════════════════════════════
# EXCEL GENERATION
# ═══════════════════════════════════════════════════════════════════

def make_excel(consultant=True):
    wb = Workbook()

    # Styles
    hdr_font = Font(name="Calibri", bold=True, color=WHITE, size=11)
    hdr_fill = PatternFill("solid", fgColor=PAN_NAVY)
    hdr_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
    hdr_border = Border(
        bottom=Side(style="thin", color=BORDER_GRAY),
        right=Side(style="thin", color=BORDER_GRAY),
    )

    gold_font = Font(name="Calibri", bold=True, color=PAN_NAVY, size=11)
    gold_fill = PatternFill("solid", fgColor=PAN_GOLD)

    cell_font = Font(name="Calibri", size=10)
    cell_align = Alignment(vertical="top", wrap_text=True)
    cell_border = Border(
        bottom=Side(style="thin", color=BORDER_GRAY),
        right=Side(style="thin", color=BORDER_GRAY),
    )

    fill_yellow = PatternFill("solid", fgColor=FILL_YELLOW)
    fill_purple = PatternFill("solid", fgColor=FILL_PURPLE)
    fill_alt = PatternFill("solid", fgColor=LIGHT_GRAY)

    req_font = Font(name="Calibri", size=10, color=RED, bold=True)
    cond_font = Font(name="Calibri", size=10, color=AMBER, bold=True)

    # ── Instructions tab ──────────────────────────────────────────
    ws = wb.active
    ws.title = "Instructions"
    ws.sheet_properties.tabColor = PAN_NAVY

    instr = [
        ("AI Red Teaming — Technical Requirements Document", ""),
        ("", ""),
        ("HOW TO USE THIS WORKBOOK", ""),
        ("", ""),
        ("This workbook collects customer-specific information needed to configure",
         "and deploy AIRS AI Red Teaming."),
        ("", ""),
        ("STRUCTURE:", ""),
        ("Part A — Engagement Context", "Asked once per customer"),
        ("Part B — Platform Configuration", "Asked once per engagement"),
        ("Part C — Per-Target Configuration",
         "Duplicate this tab for each target"),
        ("Day-2 Appendix", "Optional — address after first scan cycle"),
        ("", ""),
        ("FIELD MARKERS:", ""),
        ("*  Required", "Engagement cannot proceed without this"),
        ("†  Conditional", "Required only if a prior answer triggers it"),
        ("(unmarked)  Optional", "Improves engagement quality"),
        ("", ""),
        ("RESPONSE COLUMN:", "Fill in yellow cells with customer answers"),
        ("", ""),
        ("ADDING TARGETS:", "Right-click the 'Part C — Target 1' tab → "
         "'Move or Copy' → check 'Create a copy'. Rename the copy for each "
         "additional target."),
    ]
    if consultant:
        instr.append(("", ""))
        instr.append(("CONSULTANT GUIDANCE:", "Purple column contains internal "
                       "guidance and field tips. Remove or hide this column "
                       "before sharing with customer."))

    for i, (a, b) in enumerate(instr, 1):
        ws.cell(row=i, column=1, value=a).font = Font(
            name="Calibri", size=14 if i == 1 else 11,
            bold=(i == 1 or i in (3, 7, 13, 18, 20, 22)),
            color=PAN_NAVY if i == 1 else "333333",
        )
        ws.cell(row=i, column=2, value=b).font = Font(
            name="Calibri", size=11, color="666666")

    ws.column_dimensions["A"].width = 50
    ws.column_dimensions["B"].width = 60

    # ── Question tabs ─────────────────────────────────────────────
    for part_data in PARTS:
        ws = wb.create_sheet(title=part_data["tab_name"])
        ws.sheet_properties.tabColor = PAN_GOLD if "Day-2" in part_data["tab_name"] else PAN_NAVY

        headers = ["§", "Question", "Format", "Options",
                    "Response", "Notes"]
        if consultant:
            headers.append("Consultant Guidance")

        # Header row
        for c, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=c, value=h)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = hdr_align
            cell.border = hdr_border

        ws.freeze_panes = "A2"
        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 55
        ws.column_dimensions["C"].width = 12
        ws.column_dimensions["D"].width = 40
        ws.column_dimensions["E"].width = 40
        ws.column_dimensions["F"].width = 30
        if consultant:
            ws.column_dimensions["G"].width = 55

        row = 2
        for section in part_data["sections"]:
            # Section header row (gold)
            marker_text = f"§{section['num']}"
            title_text = f"{section['title']}  [{section['tier']}]"
            if section.get("feeds"):
                title_text += f"  →  {section['feeds']}"

            for c in range(1, len(headers) + 1):
                cell = ws.cell(row=row, column=c)
                cell.fill = gold_fill
                cell.font = gold_font
                cell.border = cell_border
            ws.cell(row=row, column=1, value=marker_text)
            ws.cell(row=row, column=2, value=title_text)
            row += 1

            # Target table (§2)
            if section.get("target_table"):
                sub_headers = ["#", "Application Name", "Type", "Use Case",
                               "Hosting", "Priority"]
                for c, h in enumerate(sub_headers, 1):
                    cell = ws.cell(row=row, column=c, value=h)
                    cell.font = Font(name="Calibri", bold=True, size=10)
                    cell.border = cell_border
                row += 1
                type_opts = '"Model,Application,Agent"'
                host_opts = '"Cloud,On-prem,Hybrid"'
                pri_opts = '"High,Medium,Low"'
                dv_type = DataValidation(type="list", formula1=type_opts)
                dv_host = DataValidation(type="list", formula1=host_opts)
                dv_pri = DataValidation(type="list", formula1=pri_opts)
                ws.add_data_validation(dv_type)
                ws.add_data_validation(dv_host)
                ws.add_data_validation(dv_pri)

                for t in range(1, 6):
                    ws.cell(row=row, column=1, value=t).font = cell_font
                    ws.cell(row=row, column=1).border = cell_border
                    for c in range(2, 7):
                        cell = ws.cell(row=row, column=c)
                        cell.fill = fill_yellow
                        cell.font = cell_font
                        cell.alignment = cell_align
                        cell.border = cell_border
                    dv_type.add(ws.cell(row=row, column=3))
                    dv_host.add(ws.cell(row=row, column=5))
                    dv_pri.add(ws.cell(row=row, column=6))
                    row += 1

            # Risk matrix (§11)
            elif section.get("risk_matrix"):
                pri_opts = '"1 (Low),2 (Medium),3 (High),N/A"'
                dv_pri = DataValidation(type="list", formula1=pri_opts)
                ws.add_data_validation(dv_pri)

                for cat_name, subcats in RISK_CATEGORIES.items():
                    # Category header
                    cell = ws.cell(row=row, column=2, value=cat_name)
                    cell.font = Font(name="Calibri", bold=True, size=10,
                                     color=PAN_NAVY)
                    cell.border = cell_border
                    row += 1

                    for sub in subcats:
                        ws.cell(row=row, column=2, value=sub).font = cell_font
                        ws.cell(row=row, column=2).border = cell_border
                        ws.cell(row=row, column=2).alignment = cell_align

                        pri_cell = ws.cell(row=row, column=5)
                        pri_cell.fill = fill_yellow
                        pri_cell.font = cell_font
                        pri_cell.border = cell_border
                        dv_pri.add(pri_cell)

                        note_cell = ws.cell(row=row, column=6)
                        note_cell.fill = fill_yellow
                        note_cell.font = cell_font
                        note_cell.border = cell_border
                        row += 1

            # Regular questions
            for qi, q in enumerate(section.get("questions", [])):
                is_alt = qi % 2 == 1

                marker = q["marker"]
                q_text = q["text"]
                if q.get("note"):
                    q_text += f"\n({q['note']})"

                ws.cell(row=row, column=1, value=marker).font = (
                    req_font if marker == "*" else
                    cond_font if marker == "†" else cell_font
                )
                ws.cell(row=row, column=1).border = cell_border
                ws.cell(row=row, column=1).alignment = Alignment(
                    horizontal="center", vertical="top")

                ws.cell(row=row, column=2, value=q_text).font = cell_font
                ws.cell(row=row, column=2).alignment = cell_align
                ws.cell(row=row, column=2).border = cell_border

                ws.cell(row=row, column=3, value=q["fmt"]).font = Font(
                    name="Consolas", size=9, color="006DCC")
                ws.cell(row=row, column=3).border = cell_border
                ws.cell(row=row, column=3).alignment = Alignment(vertical="top")

                opts_text = " / ".join(q["options"]) if q["options"] else ""
                ws.cell(row=row, column=4, value=opts_text).font = Font(
                    name="Calibri", size=9, color="666666")
                ws.cell(row=row, column=4).alignment = cell_align
                ws.cell(row=row, column=4).border = cell_border

                resp_cell = ws.cell(row=row, column=5)
                resp_cell.fill = fill_yellow
                resp_cell.font = cell_font
                resp_cell.alignment = cell_align
                resp_cell.border = cell_border

                # Add dropdown validation for select fields
                if q["fmt"] == "select" and q["options"]:
                    opts_str = ",".join(q["options"])
                    if len(opts_str) <= 255:
                        dv = DataValidation(
                            type="list", formula1=f'"{opts_str}"',
                            allow_blank=True)
                        dv.error = "Please select from the list"
                        dv.prompt = "Select an option"
                        ws.add_data_validation(dv)
                        dv.add(resp_cell)

                notes_cell = ws.cell(row=row, column=6)
                notes_cell.fill = fill_yellow
                notes_cell.font = cell_font
                notes_cell.alignment = cell_align
                notes_cell.border = cell_border

                if is_alt:
                    for c in [2, 3, 4]:
                        ws.cell(row=row, column=c).fill = fill_alt

                row += 1

            # Contacts (§1)
            if section.get("contacts"):
                sub_headers = ["", "Role", "Name", "Email", "Notes"]
                if consultant:
                    sub_headers.append("")
                for c, h in enumerate(sub_headers, 1):
                    cell = ws.cell(row=row, column=c, value=h)
                    cell.font = Font(name="Calibri", bold=True, size=10)
                    cell.border = cell_border
                row += 1
                for contact in section["contacts"]:
                    ws.cell(row=row, column=1, value=contact["marker"]).font = (
                        req_font if contact["marker"] == "*" else cell_font
                    )
                    ws.cell(row=row, column=1).border = cell_border
                    ws.cell(row=row, column=2, value=contact["role"]).font = cell_font
                    ws.cell(row=row, column=2).border = cell_border
                    for c in range(3, 6):
                        cell = ws.cell(row=row, column=c)
                        cell.fill = fill_yellow
                        cell.font = cell_font
                        cell.border = cell_border
                    row += 1

            # Consultant guidance (merged cell spanning the section)
            if consultant and section.get("consultant_note"):
                cell = ws.cell(row=row, column=2,
                               value=section["consultant_note"])
                cell.font = Font(name="Calibri", size=9, color="6D28D9",
                                 italic=True)
                cell.fill = fill_purple
                cell.alignment = cell_align
                cell.border = cell_border
                # Also fill the guidance column marker
                label = ws.cell(row=row, column=7,
                                value="← Consultant guidance")
                label.font = Font(name="Calibri", size=9, color="7C3AED",
                                  italic=True)
                label.fill = fill_purple
                row += 1

            # Blank separator row
            row += 1

    return wb


# ═══════════════════════════════════════════════════════════════════
# WORD GENERATION
# ═══════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, cells, header=False, alt=False):
    row = table.add_row()
    for i, (text, width) in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9 if not header else 10)
        run.font.name = "Calibri"
        if header:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, PAN_NAVY)
        elif alt:
            set_cell_shading(cell, LIGHT_GRAY)
    return row


def make_word(consultant=True):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    # ── Title page ────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(120)
    run = p.add_run("AIRS AI Red Teaming")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x00, 0x29, 0x4D)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Technical Requirements Document")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x00, 0x29, 0x4D)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(40)
    if consultant:
        run = p.add_run("CONSULTANT VERSION — Contains internal guidance")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
        run.font.bold = True
    else:
        run = p.add_run("Customer Intake Form")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(60)
    run = p.add_run("Customer: ___________________________________")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Date: _______________")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(80)
    run = p.add_run("Palo Alto Networks Professional Services")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── Legend ────────────────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run("How to Use This Document")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x29, 0x4D)

    legend = [
        ("*  Required", "Engagement cannot proceed without this answer."),
        ("†  Conditional", "Required only if a prior answer triggers it."),
        ("(unmarked)  Optional", "Improves engagement quality but not blocking."),
    ]
    for label, desc in legend:
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44) if "*" in label else (
            RGBColor(0xF5, 0x9E, 0x0B) if "†" in label else
            RGBColor(0x33, 0x33, 0x33))
        run = p.add_run(f" — {desc}")
        run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.space_before = Pt(8)
    run = p.add_run("Part A")
    run.font.bold = True
    run = p.add_run(" is completed once per customer. ")
    run = p.add_run("Part B")
    run.font.bold = True
    run = p.add_run(" once per engagement. ")
    run = p.add_run("Part C")
    run.font.bold = True
    run = p.add_run(" is repeated for each target or target group.")

    doc.add_page_break()

    # ── Content ───────────────────────────────────────────────────
    for part_data in PARTS:
        # Part divider
        p = doc.add_paragraph()
        p.space_before = Pt(20)
        run = p.add_run(part_data["part"])
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x29, 0x4D)

        p = doc.add_paragraph()
        run = p.add_run(part_data["desc"])
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.italic = True

        # Add a gold line (using paragraph border)
        p = doc.add_paragraph()
        p.space_after = Pt(12)

        for section in part_data["sections"]:
            # Section heading
            tier_label = f"  [{section['tier']}]"
            p = doc.add_paragraph()
            p.space_before = Pt(16)
            run = p.add_run(f"§{section['num']}  {section['title']}")
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x29, 0x4D)
            run = p.add_run(tier_label)
            run.font.size = Pt(9)
            run.font.color.rgb = (
                RGBColor(0xEF, 0x44, 0x44) if section["tier"] == "Essential" else
                RGBColor(0xF5, 0x9E, 0x0B) if section["tier"] == "Important" else
                RGBColor(0x22, 0xC5, 0x5E) if section["tier"] == "Scoping" else
                RGBColor(0x7C, 0x3A, 0xED)
            )
            run.font.bold = True

            if section.get("feeds"):
                p = doc.add_paragraph()
                run = p.add_run(f"Feeds: {section['feeds']}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x00, 0x6D, 0xCC)
                run.font.italic = True

            # Target table (§2)
            if section.get("target_table"):
                table = doc.add_table(rows=1, cols=6)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                table.style = "Table Grid"
                headers = ["#", "Application Name", "Type", "Use Case",
                           "Hosting", "Priority"]
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(h)
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    set_cell_shading(cell, PAN_NAVY)

                for t in range(1, 6):
                    row = table.add_row()
                    prefills = {
                        0: str(t), 2: "Model / Application / Agent",
                        4: "Cloud / On-prem / Hybrid", 5: "High / Medium / Low",
                    }
                    for c in range(6):
                        cell = row.cells[c]
                        cell.text = ""
                        txt = prefills.get(c, "")
                        run = cell.paragraphs[0].add_run(txt)
                        run.font.size = Pt(9)
                        if c in (2, 4, 5):
                            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                            run.font.italic = True

            # Risk matrix (§11)
            elif section.get("risk_matrix"):
                for cat_name, subcats in RISK_CATEGORIES.items():
                    p = doc.add_paragraph()
                    p.space_before = Pt(8)
                    run = p.add_run(cat_name)
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x00, 0x29, 0x4D)

                    table = doc.add_table(rows=1, cols=3)
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    table.style = "Table Grid"
                    for i, h in enumerate(["Subcategory", "Priority", "Notes"]):
                        cell = table.rows[0].cells[i]
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(h)
                        run.font.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        set_cell_shading(cell, PAN_NAVY)

                    for si, sub in enumerate(subcats):
                        row = table.add_row()
                        row.cells[0].text = sub
                        row.cells[1].text = "1 (Low) / 2 (Medium) / 3 (High) / N/A"
                        for c in range(3):
                            for p in row.cells[c].paragraphs:
                                for r in p.runs:
                                    r.font.size = Pt(9)
                                    if c == 1:
                                        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                                        r.font.italic = True
                            if si % 2 == 1:
                                set_cell_shading(row.cells[c], LIGHT_GRAY)

            # Regular question table
            if section.get("questions"):
                table = doc.add_table(rows=1, cols=2)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                table.style = "Table Grid"

                for i, h in enumerate(["Question", "Response / Notes"]):
                    cell = table.rows[0].cells[i]
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(h)
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    set_cell_shading(cell, PAN_NAVY)

                for qi, q in enumerate(section["questions"]):
                    row = table.add_row()
                    # Question cell
                    p = row.cells[0].paragraphs[0]
                    if q["marker"]:
                        run = p.add_run(q["marker"] + " ")
                        run.font.color.rgb = (
                            RGBColor(0xEF, 0x44, 0x44) if q["marker"] == "*"
                            else RGBColor(0xF5, 0x9E, 0x0B))
                        run.font.bold = True
                        run.font.size = Pt(9)
                    run = p.add_run(q["text"])
                    run.font.size = Pt(9)

                    if q.get("note"):
                        run = p.add_run(f"\n({q['note']})")
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                        run.font.italic = True

                    # Format badge
                    run = p.add_run(f"  [{q['fmt']}]")
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x00, 0x6D, 0xCC)

                    # Response cell
                    resp = row.cells[1].paragraphs[0]
                    if q["options"]:
                        run = resp.add_run(" / ".join(q["options"]))
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                        run.font.italic = True
                    else:
                        resp.text = ""

                    # Alternate row shading
                    if qi % 2 == 1:
                        set_cell_shading(row.cells[0], LIGHT_GRAY)
                        set_cell_shading(row.cells[1], LIGHT_GRAY)

            # Contact table (§1)
            if section.get("contacts"):
                p = doc.add_paragraph()
                p.space_before = Pt(8)
                run = p.add_run("Stakeholder Contacts")
                run.font.size = Pt(11)
                run.font.bold = True

                table = doc.add_table(rows=1, cols=4)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                table.style = "Table Grid"
                for i, h in enumerate(["Role", "Name", "Email", "Notes"]):
                    cell = table.rows[0].cells[i]
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(h)
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    set_cell_shading(cell, PAN_NAVY)

                for contact in section["contacts"]:
                    row = table.add_row()
                    p = row.cells[0].paragraphs[0]
                    if contact["marker"]:
                        run = p.add_run(contact["marker"] + " ")
                        run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
                        run.font.bold = True
                        run.font.size = Pt(9)
                    run = p.add_run(contact["role"])
                    run.font.size = Pt(9)

            # Consultant notes
            if consultant and section.get("consultant_note"):
                p = doc.add_paragraph()
                p.space_before = Pt(8)

                # Purple border box effect via paragraph shading
                pPr = p._p.get_or_add_pPr()
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="F3E8FF" w:val="clear"/>')
                pPr.append(shading)

                # Left border
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'<w:left w:val="single" w:sz="12" w:space="4" w:color="7C3AED"/>'
                    f'</w:pBdr>')
                pPr.append(pBdr)

                run = p.add_run("Consultant Notes: ")
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)

                run = p.add_run(section["consultant_note"])
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x4C, 0x1D, 0x95)
                run.font.italic = True

        # Page break between parts
        doc.add_page_break()

    return doc


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════

def main():
    os.makedirs(OUT_DIR, exist_ok=True)

    files = {
        "trd-red-teaming-consultant.xlsx": ("Excel consultant", lambda: make_excel(True)),
        "trd-red-teaming-customer.xlsx": ("Excel customer", lambda: make_excel(False)),
        "trd-red-teaming-consultant.docx": ("Word consultant", lambda: make_word(True)),
        "trd-red-teaming-customer.docx": ("Word customer", lambda: make_word(False)),
    }

    for filename, (label, gen_fn) in files.items():
        path = os.path.join(OUT_DIR, filename)
        print(f"  Generating {label}... ", end="", flush=True)
        result = gen_fn()
        if filename.endswith(".xlsx"):
            result.save(path)
        else:
            result.save(path)
        print(f"✓  {path}")

    print(f"\nDone — 4 files in {OUT_DIR}/")


if __name__ == "__main__":
    main()
